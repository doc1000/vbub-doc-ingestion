"""Tests for DocxExtractor."""

import io
from pathlib import Path

import docx
import pytest

from vbub_doc_ingestion.extractors.docx_extractor import DocxExtractor

FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture
def extractor() -> DocxExtractor:
    return DocxExtractor()


@pytest.fixture
def sample_docx_bytes() -> bytes:
    return (FIXTURES / "sample.docx").read_bytes()


def test_docx_returns_paragraph_text(extractor: DocxExtractor, sample_docx_bytes: bytes) -> None:
    result = extractor.extract(sample_docx_bytes, "sample.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "first body paragraph" in result.clean_text
    assert "second body paragraph" in result.clean_text


def test_docx_title_from_core_properties(extractor: DocxExtractor, sample_docx_bytes: bytes) -> None:
    result = extractor.extract(sample_docx_bytes, "sample.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.title == "Sample DOCX Document"


def test_docx_title_falls_back_to_heading(extractor: DocxExtractor) -> None:
    document = docx.Document()
    # No core_properties.title set
    document.add_heading("Heading Level One", level=1)
    document.add_paragraph("Some body text here.")
    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    result = extractor.extract(docx_bytes, "nocoretitle.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.title == "Heading Level One"


def test_docx_title_falls_back_to_first_paragraph(extractor: DocxExtractor) -> None:
    document = docx.Document()
    document.add_paragraph("First paragraph as fallback title.")
    document.add_paragraph("Second paragraph body text.")
    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    result = extractor.extract(docx_bytes, "nohdr.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.title == "First paragraph as fallback title."


def test_docx_empty_paragraphs_are_skipped(extractor: DocxExtractor) -> None:
    document = docx.Document()
    document.add_paragraph("")
    document.add_paragraph("   ")
    document.add_paragraph("Real content here.")
    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    result = extractor.extract(docx_bytes, "empty_paras.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.clean_text.strip() == "Real content here."


def test_docx_parser_name_and_version(extractor: DocxExtractor, sample_docx_bytes: bytes) -> None:
    result = extractor.extract(sample_docx_bytes, "sample.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.parser_name == "DocxExtractor"
    assert result.parser_version == "0.1.0"


def test_docx_warnings_empty(extractor: DocxExtractor, sample_docx_bytes: bytes) -> None:
    result = extractor.extract(sample_docx_bytes, "sample.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.warnings == []
