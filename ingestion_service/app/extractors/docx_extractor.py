"""DOCX extractor using python-docx.

Reads all paragraphs from the document body, skips empty ones,
and joins them with single newlines. Returns raw (un-normalised) text;
normalisation is applied by a downstream step.

Title resolution order:
  1. core_properties.title (document property, if non-empty)
  2. First paragraph with a Heading style (e.g. "Heading 1")
  3. First non-empty paragraph text
  4. None
"""

import io
from typing import Optional

import docx
from docx.document import Document as DocxDocument

from ingestion_service.app.domain.contracts import ExtractionResult

_PARSER_NAME = "DocxExtractor"
_PARSER_VERSION = "0.1.0"


def _resolve_title(doc: DocxDocument, paragraphs: list[str]) -> Optional[str]:
    """Derive a document title from core properties, headings, or first paragraph."""
    # 1. Document core property
    core_title = (doc.core_properties.title or "").strip()
    if core_title:
        return core_title

    # 2. First heading-style paragraph
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            text = para.text.strip()
            if text:
                return text

    # 3. First non-empty paragraph
    for text in paragraphs:
        if text.strip():
            return text.strip()

    return None


class DocxExtractor:
    """Extracts text from DOCX files using python-docx.

    Implements the DocumentExtractor Protocol.
    """

    def extract(
        self,
        file_bytes: bytes,
        filename: str,
        canonical_mime: str,
    ) -> ExtractionResult:
        """Extract paragraph text from a DOCX binary.

        Empty paragraphs are skipped. Paragraphs are joined with newlines.

        Args:
            file_bytes:     raw bytes of the DOCX file.
            filename:       original filename (used for error context only).
            canonical_mime: not used for dispatch here.

        Returns:
            ExtractionResult with joined paragraph text and optional title.
        """
        doc = docx.Document(io.BytesIO(file_bytes))

        texts = [para.text for para in doc.paragraphs if para.text.strip()]
        body = "\n".join(texts)
        title = _resolve_title(doc, texts)

        return ExtractionResult(
            parser_name=_PARSER_NAME,
            parser_version=_PARSER_VERSION,
            title=title,
            clean_text=body,
            warnings=[],
        )
