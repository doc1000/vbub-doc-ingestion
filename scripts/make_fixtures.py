"""One-off script to generate test fixtures for Phase 4 and Phase 5.

Creates:
  tests/fixtures/sample.pdf  — 3-page PDF with repeated header/footer
  tests/fixtures/sample.docx — DOCX with core title, heading, and body paragraphs
  tests/fixtures/sample.xlsx — 2-sheet XLSX: text-rich sheet + numeric-only sheet

Run once from the project root: python scripts/make_fixtures.py
"""

import fitz  # PyMuPDF
import docx
import openpyxl

REPEATED_HEADER = "Acme Corp Confidential"
REPEATED_FOOTER = "Internal Use Only"


def make_sample_pdf(path: str) -> None:
    doc = fitz.open()
    for i in range(1, 4):
        lines = [
            REPEATED_HEADER,
            "",
            f"Chapter {i}: Body Content",
            "",
            f"This is substantive body text on page {i}.",
            "It contains useful information for downstream processing.",
            "",
            REPEATED_FOOTER,
            str(i),
        ]
        page = doc.new_page()
        page.insert_text((72, 72), "\n".join(lines), fontsize=11)
    doc.set_metadata({"title": "Quarterly Research Report"})
    doc.save(path)
    doc.close()
    print(f"Written: {path}")


def make_sample_docx(path: str) -> None:
    document = docx.Document()
    document.core_properties.title = "Sample DOCX Document"
    document.add_heading("Introduction", level=1)
    document.add_paragraph(
        "This is the first body paragraph of the sample DOCX document."
    )
    document.add_paragraph(
        "This is the second body paragraph with additional content for testing."
    )
    document.save(path)
    print(f"Written: {path}")


def make_sample_xlsx(path: str) -> None:
    wb = openpyxl.Workbook()

    # Sheet 1: text-rich data
    ws1 = wb.active
    ws1.title = "Research Topics"
    ws1.append(["Topic", "Author", "Summary"])
    ws1.append(["Machine Learning", "Alice Smith", "Overview of supervised learning methods"])
    ws1.append(["Natural Language Processing", "Bob Jones", "Text classification and embeddings"])
    ws1.append([1.0, 2.5, 3.7])   # numeric-only row — should be suppressed
    ws1.append([4.2, 5.1, 6.8])   # numeric-only row — should be suppressed
    ws1.append(["Computer Vision", "Carol White", "Image recognition and object detection"])

    # Sheet 2: all numeric — all body rows should be suppressed (only header kept)
    ws2 = wb.create_sheet(title="Raw Numbers")
    ws2.append(["Val1", "Val2", "Val3"])
    ws2.append([10.0, 20.0, 30.0])
    ws2.append([40.0, 50.0, 60.0])

    wb.save(path)
    print(f"Written: {path}")


if __name__ == "__main__":
    make_sample_pdf("tests/fixtures/sample.pdf")
    make_sample_docx("tests/fixtures/sample.docx")
    make_sample_xlsx("tests/fixtures/sample.xlsx")
